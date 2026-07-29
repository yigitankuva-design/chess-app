from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import os

LEVELS = [
    {
        "no": "1",
        "baslik": "TEMEL DÜZEY",
        "emoji": "🌱",
        "renk": RGBColor(0x27, 0xAE, 0x60),
        "hex": "27AE60",
        "dersler": [
            "Satranç Tahtası ve Koordinatlar",
            "Taşları Tanıyalım",
            "Piyon Nasıl Hareket Eder?",
            "At Nasıl Hareket Eder?",
            "Fil Nasıl Hareket Eder?",
            "Kale Nasıl Hareket Eder?",
            "Vezir Nasıl Hareket Eder?",
            "Şah Nasıl Hareket Eder?",
            "Rok Hamlesi",
            "Şah Çekme ve Mat Nedir?",
        ],
    },
    {
        "no": "2",
        "baslik": "BAŞLANGIÇ DÜZEYİ",
        "emoji": "😊",
        "renk": RGBColor(0x29, 0x80, 0xB9),
        "hex": "2980B9",
        "dersler": [
            "Taşların Değeri",
            "Açılışta Merkezi Kontrol Et",
            "Taşlarını Geliştir",
            "Şahını Güvene Al",
            "Çatal Taktiği",
            "İğne Taktiği",
            "Bağlama Taktiği",
            "Çekme ve Sapıştırma",
            "Kral + Vezir ile Mat",
            "Kral + Kale ile Mat",
        ],
    },
    {
        "no": "3",
        "baslik": "ORTA DÜZEY",
        "emoji": "😎",
        "renk": RGBColor(0x8E, 0x44, 0xAD),
        "hex": "8E44AD",
        "dersler": [
            "e4 Açılışları",
            "d4 Açılışları",
            "Piyon Yapısı ve Zayıf Kareler",
            "Açık Hatlar ve Sütunlar",
            "İki Fil Avantajı",
            "Orta Oyun Planlaması",
            "Taktik Kombinasyonlar — Kurban",
            "Piyon Son Oyunu Temelleri",
            "Kale Son Oyunu Temelleri",
            "Pozisyonel Satranç",
        ],
    },
    {
        "no": "4",
        "baslik": "İLERİ DÜZEY",
        "emoji": "🔥",
        "renk": RGBColor(0xC0, 0x39, 0x2B),
        "hex": "C0392B",
        "dersler": [
            "Açılış Teorisi ve Repertuar",
            "Stratejik Kurban",
            "Dinamik Oyun ve Girişim",
            "Karmaşık Taktik Hesaplama",
            "Profülaktik Düşünme",
            "Zayıflıklara Baskı Yapma",
            "Lucena Pozisyonu",
            "Philidor Pozisyonu",
            "Turnuva Psikolojisi",
            "Rakip Analizi ve Hazırlık",
        ],
    },
]


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_cell_border(cell, **kwargs):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "left", "bottom", "right"):
        tag = OxmlElement(f"w:{side}")
        tag.set(qn("w:val"), "single")
        tag.set(qn("w:sz"), "6")
        tag.set(qn("w:space"), "0")
        tag.set(qn("w:color"), kwargs.get(side, "DDDDDD"))
        tcBorders.append(tag)
    tcPr.append(tcBorders)


doc = Document()

# Sayfa kenar boşlukları
section = doc.sections[0]
section.top_margin    = Cm(2)
section.bottom_margin = Cm(2)
section.left_margin   = Cm(2.5)
section.right_margin  = Cm(2.5)

# ── Başlık ──────────────────────────────────────
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run("BEA AKADEMİ GELİŞİM SİSTEMİ")
run.bold = True
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x2E)

sub = doc.add_paragraph()
sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
sub_run = sub.add_run("Ders Fasikülü — Düzey ve İçerik Tablosu")
sub_run.italic = True
sub_run.font.size = Pt(11)
sub_run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph()  # boşluk

# ── Her düzey için tablo ─────────────────────────
for lv in LEVELS:
    color = lv["renk"]
    hex_c = lv["hex"]

    # Düzey başlığı
    hdr = doc.add_paragraph()
    hdr.alignment = WD_ALIGN_PARAGRAPH.LEFT
    hdr_run = hdr.add_run(f"  {lv['emoji']}  Fasikül {lv['no']} — {lv['baslik']}")
    hdr_run.bold = True
    hdr_run.font.size = Pt(13)
    hdr_run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Başlık paragrafına renkli arka plan (shading via XML)
    pPr = hdr._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_c)
    pPr.append(shd)

    # Tablo: 3 sütun — No | Ders Adı | Notlar
    tbl = doc.add_table(rows=1, cols=3)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

    # Sütun genişlikleri
    for i, width in enumerate([Cm(1.8), Cm(9.5), Cm(5.0)]):
        for cell in tbl.columns[i].cells:
            cell.width = width

    # Tablo başlık satırı
    hrow = tbl.rows[0]
    headers = ["No", "Ders Adı", "Notlar / Süre"]
    for i, h in enumerate(headers):
        cell = hrow.cells[i]
        set_cell_bg(cell, hex_c)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(h)
        run.bold = True
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

    # Ders satırları
    for idx, ders in enumerate(lv["dersler"], start=1):
        row = tbl.add_row()
        bg = "FFFFFF" if idx % 2 == 0 else "F7F7F7"

        # No
        c0 = row.cells[0]
        set_cell_bg(c0, bg)
        c0.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p0 = c0.paragraphs[0]
        p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r0 = p0.add_run(str(idx))
        r0.font.size = Pt(10)
        r0.bold = True
        r0.font.color.rgb = color

        # Ders adı
        c1 = row.cells[1]
        set_cell_bg(c1, bg)
        c1.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p1 = c1.paragraphs[0]
        p1.paragraph_format.left_indent = Cm(0.3)
        r1 = p1.add_run(ders)
        r1.font.size = Pt(10)

        # Notlar (boş, doldurulacak)
        c2 = row.cells[2]
        set_cell_bg(c2, bg)
        r2 = c2.paragraphs[0].add_run("")
        r2.font.size = Pt(10)

    doc.add_paragraph()  # düzeyler arası boşluk

# ── Alt not ─────────────────────────────────────
footer_p = doc.add_paragraph()
footer_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
footer_run = footer_p.add_run("BEA Akademi Gelişim Sistemi  |  Satranç Müfredatı")
footer_run.font.size = Pt(9)
footer_run.italic = True
footer_run.font.color.rgb = RGBColor(0xAA, 0xAA, 0xAA)

# Kaydet
desktop = os.path.join(os.path.expanduser("~"), "Desktop")
path = os.path.join(desktop, "BEA_Akademi_Ders_Fasikulu.docx")
doc.save(path)
print(f"Kaydedildi: {path}")
